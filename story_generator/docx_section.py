# modules/story_generator/docx_sections.py
import requests
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .docx_helper import _make_cell_bold, _merge_cells

class DocxSections:


    def __init__(self, story_generator):
        self.story_generator = story_generator

    def add_user_story_section(self, doc: Document, clickup_data: Dict):
        """Third Page - User Story with Preconditions (exact format)"""
        title = doc.add_heading('User Story', 1)
        
        # Generate concise user story from ClickUp data using OpenAI
        user_story_text = self._generate_concise_user_story(clickup_data)
        
        # Add user story in 3 separate bold lines
        lines = user_story_text.split('\n')
        for line in lines:
            if line.strip():
                para = doc.add_paragraph()
                run = para.add_run(line.strip())
                run.bold = True
        
        doc.add_paragraph("---")
        
        # Preconditions table (exact format)
        preconditions_title = doc.add_heading('Preconditions', 2)
        preconditions_table = doc.add_table(rows=4, cols=2)
        preconditions_table.style = 'Table Grid'
        
        # Get screen name from ClickUp data
        screen_name = clickup_data.get("title", "[SCREEN_NAME]")
        
        preconditions_data = [
            ("", "Precondition"),
            ("1", "User has logged in to the System as Tenant or Tenant User"),
            ("2", "User's role has permissions associated with this screen"), 
            ("3", f"User has navigated to {screen_name}")
        ]
        
        for i, (num, condition) in enumerate(preconditions_data):
            preconditions_table.cell(i, 0).text = str(num)
            preconditions_table.cell(i, 1).text = condition
        
        # Add Figma/Wireframe Link section
        doc.add_heading('Figma/Wireframe Link:', 2)
        figma_link = clickup_data.get("figma_link", "ADD LINK TO ARTICLE DESCRIBING RELEVANT UI")
        doc.add_paragraph(figma_link)

    def _generate_concise_user_story(self, clickup_data: Dict) -> str:
            """Generate a human-readable, concise user story using OpenAI"""
            title = clickup_data.get("title", "")
            description = clickup_data.get("description_part1", "")
            business_case = clickup_data.get("business_case", "")

            prompt = f"""
        Create a concise 3-line user story in the format:
        As a [role]
        I want [feature/action]
        So that [benefit/outcome]
        
        Based on this information:
        Title: {title}
        Description: {description}
        Business Case: {business_case}
        
        Keep it to exactly 3 lines, each line should be clear and concise.
        Return only the 3 lines without any additional text.
        """

            try:
                response = self.story_generator.client.chat.completions.create(
                    model=self.story_generator.config.deployment_name,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=350,
                    temperature=0.4  # Lowered for more consistency
                )
                user_story = response.choices[0].message.content.strip()

                # Clean up the response
                lines = [l.strip() for l in user_story.split('\n') if l.strip()]
                
                # Remove any numbering or bullet points
                cleaned_lines = []
                for line in lines:
                    # Remove common prefixes
                    line = line.lstrip('123456789.-*• ')
                    if line.lower().startswith(('as a', 'i want', 'so that')):
                        cleaned_lines.append(line)
                
                # Ensure exactly 3 lines
                if len(cleaned_lines) >= 3:
                    return '\n'.join(cleaned_lines[:3])
                else:
                    # Better fallback using actual context
                    return self._create_fallback_story(title, description, business_case)

            except Exception as e:
                print(f"Error generating user story with OpenAI: {e}")
                return self._create_fallback_story(title, description, business_case)
            
    # ---------------- Acceptance Criteria + Business Rules ----------------
    def add_acceptance_criteria_table(self, doc: Document, figma_data: Dict, clickup_data: Dict):
            """Fourth Page - Acceptance Criteria Table (exact format from your reference)"""
            title = doc.add_heading('Acceptance Criteria', 1)
            
            screens = figma_data.get("screens", [])
            
            # Calculate total rows: headers (3) + screens + step/heading (1) + permissions (2)
            total_rows = 3 + len(screens) + 1 + 2
            table = doc.add_table(rows=total_rows, cols=5)
            table.style = 'Table Grid'
            
            # Headers (exact format from your reference image)
            headers_data = [
                ["ID", "", "", "", ""],
                ["Name", "", "", "", ""],
                ["Business Use Case", "", "", "", ""],
                ["Module/Submodule/Feature", "Current Screen", "Figma Screenshot", "Business Rules", "Comments"]
            ]
            
         
            for row_idx, headers in enumerate(headers_data):
                for col_idx, header in enumerate(headers):
                    table.cell(row_idx, col_idx).text = header
                    if row_idx == 3:  # Make sub-headers bold
                        _make_cell_bold(table.cell(row_idx, col_idx))
        
            story_name = clickup_data.get("title", "")
            business_use_case = clickup_data.get("business_case", "")
            module_info = clickup_data.get("module", "")
            
     
            for i, screen in enumerate(screens):
                row_idx = i + 4  # Start from row 5 (after headers)
                if row_idx >= len(table.rows):
                    break
                
 
                step_heading = self.story_generator.generate_step_heading(screen.get('frame_summary', ''))
                table.cell(row_idx, 0).text = f"{i + 1} {step_heading}"

          
                frame_url = screen.get('frame_url', '')
                cell_img = table.cell(row_idx, 2)
                try:
                    response = requests.get(frame_url, timeout=10)
                    if response.status_code == 200:
                        image_bytes = BytesIO(response.content)
                        paragraph = cell_img.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(image_bytes, width=Inches(2))
                    else:
                        cell_img.text = f"(Image unavailable: {response.status_code})"
                except Exception as e:
                    print(f" Error loading frame image: {e}")
                    cell_img.text = "(Failed to load image)"

              
                business_rules = self._generate_business_rules_from_screen(screen, clickup_data)
                cell = table.cell(row_idx, 3)

           
                for paragraph in cell.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)

                lines = [l.strip() for l in business_rules.split("\n") if l.strip()]
                interactions = screen.get("interactions", [])

                for j, line in enumerate(lines):
                    para = cell.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if j > 0 and (j - 1) < len(interactions):
                        to_url = interactions[j - 1].get("to_url", "")
                        if to_url:
                            print(f" Adding to_url image: {to_url}")
                            try:
                                response = requests.get(to_url, timeout=30)
                                if response.status_code == 200:
                                    image_bytes = BytesIO(response.content)
                                    img_para = cell.add_paragraph()
                                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = img_para.add_run()
                                    run.add_picture(image_bytes, width=Inches(2))
                                else:
                                    cell.add_paragraph(f"(Image unavailable: {response.status_code})")
                            except Exception as e:
                                print(f"⚠️ Error adding image from {to_url}: {e}")
                                cell.add_paragraph(f"(Failed to load image: {to_url})")

                cell.add_paragraph("")  # Add a small space at the end
                table.cell(row_idx, 4).text = ""

            last_content_row = 4 + len(screens)
            if last_content_row < len(table.rows):
                table.cell(last_content_row, 0).text = "Permissions"
                _merge_cells(table, last_content_row, 0, last_content_row, 4)
            
            if last_content_row + 1 < len(table.rows):
                table.cell(last_content_row + 1, 0).text = "Dashboard Notifications"
                _merge_cells(table, last_content_row + 1, 0, last_content_row + 1, 4)

    
    def _generate_business_rules_from_screen(self, screen: Dict, clickup_data: Dict) -> str:
        """
        Generate rich business rules combining ClickUp narrative + Figma flow.
        Writes exactly (1 + number_of_interactions) points in natural language.
        """

        frame_summary = screen.get("frame_summary", "")
        interactions = screen.get("interactions", [])
        clickup_desc = clickup_data.get("description_part1", "")
        business_case = clickup_data.get("business_case", "")
        comments = " ".join([c.get("content", "") for c in clickup_data.get("comments", [])])

        # Combine ClickUp context into one narrative
        full_clickup_context = (
                f"Description:\n{clickup_desc}\n\n"
                f"Business Case:\n{business_case}\n\n"
                f"Comments:\n{comments}\n"
        )

        if not frame_summary and not interactions:
            return "• Business rules will be defined based on screen functionality."

        # Build readable interaction map
        interactions_text = ""
        for i, inter in enumerate(interactions):
            from_desc = inter.get("from_summary", "")
            to_desc = inter.get("to_summary", "")
            if from_desc or to_desc:
                interactions_text += f"\nInteraction {i+1}:\nFrom: {from_desc}\nTo: {to_desc}\n"

        # Refined prompt
        prompt = f"""
            Based on this screen description and ALL interactions, generate business rules as bullet points.
            
            SCREEN OVERVIEW:
            {frame_summary}
            
            ALL INTERACTIONS:
            {interactions_text}

            Clickup:
            {full_clickup_context}
            
            Generate exactly:
            - 1 bullet point for the main screen purpose/functionality
            - Then 1 bullet point for EACH interaction (from_summary + to_summary combined)
            
            Total bullet points should be: {1+len(interactions)}
            
            Format each bullet point as:
            • [Clear business rule describing the functionality or interaction]
            
            Keep each bullet point concise (25-35 words maximum).
            Focus on business functionality, user workflows, and system behavior.
            
            Return ONLY the bullet points without any additional text.
            """

        try:
            response = self.story_generator.client.chat.completions.create(
                model=self.story_generator.config.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=700,
                temperature=0.35
            )
            business_rules = response.choices[0].message.content.strip()
            return business_rules

        except Exception as e:
            print(f" GPT generation failed: {e}")
            return self._create_fallback_business_rules(frame_summary, interactions)

    def _create_fallback_business_rules(self, frame_summary: str, interactions: List[Dict]) -> str:
            """Create business rules from actual data when OpenAI fails"""
            rules = []
            
            # Add main screen rule from frame_summary
            if frame_summary:
                # Extract the main purpose from frame_summary
                purpose_line = ""
                lines = frame_summary.split('\n')
                for line in lines:
                    if "purpose" in line.lower() or "purpose:" in line.lower():
                        purpose_line = line
                        break
                
                if purpose_line:
                    # Clean up the purpose line
                    purpose = purpose_line.split(':', 1)[-1].strip()
                    purpose = purpose.replace('**', '').replace('*', '')
                    rules.append(f"• {purpose}")
                else:
                    # Use first line of frame_summary
                    first_line = frame_summary.split('\n')[0].strip()
                    first_line = first_line.replace('**', '').replace('*', '')
                    rules.append(f"• {first_line}")
            
            # Add rules for each interaction
            for i, interaction in enumerate(interactions):
                from_desc = interaction.get('from_summary', '')
                to_desc = interaction.get('to_summary', '')
                
                if from_desc and to_desc:
                    # Extract action from from_summary
                    action = ""
                    from_lines = from_desc.split('\n')
                    for line in from_lines:
                        if "click" in line.lower() or "select" in line.lower() or "navigate" in line.lower():
                            action = line.strip()
                            break
                    if not action:
                        action = from_lines[0].strip() if from_lines else "User interaction"
                    
                    # Extract destination from to_summary
                    destination = ""
                    to_lines = to_desc.split('\n')
                    for line in to_lines:
                        if "purpose" in line.lower() or "screen" in line.lower() or "navigated" in line.lower():
                            destination = line.strip()
                            break
                    if not destination:
                        destination = to_lines[0].strip() if to_lines else "destination screen"
                    
                    # Clean up the texts
                    action = action.replace('**', '').replace('*', '').replace('1.', '').replace('2.', '').replace('3.', '')
                    destination = destination.replace('**', '').replace('*', '').replace('1.', '').replace('2.', '').replace('3.', '')
                    
                    rules.append(f"• {action} leads to {destination}")
            
            # If no rules were created, provide generic ones
            if not rules:
                rules = [
                    "• Users can interact with screen elements and components",
                    "• System provides navigation between different screens and features",
                    "• Business functionality supports user workflows and processes"
                ]
            
            return "\n".join(rules)
    # ---------------- Reference Requirements & Checklists ----------------
    def add_reference_requirements(self, doc: Document, clickup_data: Dict):
        """Fifth Page - Reference Requirements"""
        doc.add_heading('Reference Requirements', 1)
        
        # Get data from ClickUp
        data_section = clickup_data.get("data_requirements", "User data, System configuration")
        mobile_section = clickup_data.get("mobile_scope", "Included in above requirements")
        permissions_section = clickup_data.get("permissions_settings", "Standard permissions apply")
        
        doc.add_paragraph(f"1. Data: {data_section}")
        doc.add_paragraph(f"2. Mobile: {mobile_section}")
        doc.add_paragraph(f"3. Permissions/Notification Settings: {permissions_section}")

    def add_ba_cross_check(self, doc: Document):
        """Sixth Page - BA Cross-Check List (exact format)"""
        doc.add_heading('BA Cross-Check List', 1)
        
        table = doc.add_table(rows=16, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        table.cell(0, 0).text = "Reviewer"
        table.cell(1, 0).text = "Review Date"
        
        # Sub-headers
        sub_headers = ["SL No.", "Condition", "Covered/Yes/No/NA", "Comments"]
        for i, header in enumerate(sub_headers):
            table.cell(2, i).text = header
            _make_cell_bold(table.cell(2, i))
        
        # Conditions (exact list from reference)
        conditions = [
            "Are all flow with Business Rules covered?",
            "Is Figma as per standard?",
            "Are all attached Screenshots as per latest Figma?",
            "Is permission section Covered",
            "Is Dashboard Notification Covered?",
            "Is Email Notification Covered? With respective email link navigation.",
            "Is all other functionality impact covered?",
            "Is history impact covered?",
            "Mobile Scope is Defined?",
            "Admin/Super Admin impact are covered?",
            "Is Correct Figma Link attached in the story?",
            "Is Story in Published Mode?",
            "Is Published Story link attached in the Click up?"
        ]
        
        for i, condition in enumerate(conditions):
            row_idx = i + 3
            if row_idx < len(table.rows):
                table.cell(row_idx, 0).text = str(i + 1)
                table.cell(row_idx, 1).text = condition

    def add_requirements_acceptance(self, doc: Document):
        """Seventh Page - Requirements Acceptance By Product Owner"""
        doc.add_heading('Requirements Acceptance By Product Owner', 1)
        
        table = doc.add_table(rows=12, cols=3)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = "Reviewer"
        table.cell(1, 0).text = "Review Date"
        
        sub_headers = ["Condition", "Covered(Yes/No/NA)", "Comments"]
        for i, header in enumerate(sub_headers):
            table.cell(2, i).text = header
            _make_cell_bold(table.cell(2, i))
        
        conditions = [
            "Functional Flow and Dependencies Covered?",
            "Figma covered for all the scenarios for Web and Mobile both?",
            "UI Elements explained?",
            "Validations covered?",
            "Decisions covered for existing data / flow ? (if existing flow / data behavior is going to be impacted)",
            "Admin / Super Admin touchpoints covered?",
            "Roles / Rights / Permissions covered?",
            "Email / Notification content covered?",
            "Mobile Scope defined?"
        ]
        
        for i, condition in enumerate(conditions):
            row_idx = i + 3
            if row_idx < len(table.rows):
                table.cell(row_idx, 0).text = condition
