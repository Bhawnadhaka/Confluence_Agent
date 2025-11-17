# modules/story_generator/confluence_agent.py
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any
from docx import Document

from .docx_helper import _make_cell_bold
from story_generator.docx_section import DocxSections


class ConfluenceAgent:
    def __init__(self, story_generator, base_path: str = None):
        self.story_generator = story_generator
        self.base_path = base_path or os.getcwd()
        self.docx_sections = DocxSections(self.story_generator)

    def generate_complete_story(
        self,
        clickup_data: Dict[str, Any],
        figma_data: Dict[str, Any],
    ) -> Document:
        """
        Generate complete Confluence story as a Word document only.
        """
        if not clickup_data or not figma_data:
            raise ValueError("Missing in-memory data for ClickUp or Figma summaries.")

        data = {"clickup": clickup_data, "figma": figma_data}
        return self._generate_word_document(data)

    def _generate_word_document(self, data: Dict[str, Any]) -> Document:
        """Generate structured Word story document."""
        clickup_data = data.get("clickup", {})
        figma_data = data.get("figma", {"screens": []})
        doc = Document()

        # Page 1: Page Properties
        self._add_page_properties(doc, clickup_data)

        # Page 2: Table of Contents
        self._add_table_of_contents(doc)

        # Page 3: User Story + Preconditions
        self.docx_sections.add_user_story_section(doc, clickup_data)

        # Page 4: Acceptance Criteria
        self.docx_sections.add_acceptance_criteria_table(doc, figma_data, clickup_data)

        # Page 5: Reference Requirements
        doc.add_page_break()
        self.docx_sections.add_reference_requirements(doc, clickup_data)

        # Page 6: BA Cross Check
        doc.add_page_break()
        self.docx_sections.add_ba_cross_check(doc)

        # Page 7: Requirements Acceptance
        doc.add_page_break()
        self.docx_sections.add_requirements_acceptance(doc)

        return doc

    def _add_page_properties(self, doc: Document, clickup_data: Dict[str, Any]):
        title = doc.add_heading("Page Properties", 1)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"

        assignees = clickup_data.get("assignees", "")
        responsible_text = f"{assignees}" if assignees else "@... MENTION RESPONSIBLE PERSON HERE"

        properties = [
            ("User Story", "ADD LINK TO STORY IN AZURE DEVOPS HERE"),
            ("Epic/Feature", "ADD LINK TO RELATED EPIC IN AZURE DEVOPS HERE"),
            ("Requirement status", ""),
            ("Responsible", responsible_text),
            ("Tracking", "Changes for the current version are marked in Green"),
            ("Current Version", "Editable version TBD in expander"),
        ]

        for i, (key, value) in enumerate(properties):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
            if key in ["Responsible", "Requirement status"]:
                _make_cell_bold(table.cell(i, 1))

    def _add_table_of_contents(self, doc: Document):
        title = doc.add_heading("Table of Contents", 1)
        items = [
            "Table of Contents",
            "User Story",
            "Preconditions",
            "Figma/Wireframe Link",
            "Acceptance Criteria",
            "Reference Requirements",
            "BA Cross-Check List",
            "Requirements Acceptance By Product Owner",
        ]
        for item in items:
            doc.add_paragraph(f"• {item}", style="List Bullet")

    # ------------------------------------------------------------------
    # Saving (Word only)
    # ------------------------------------------------------------------
    def save_story_to_file(self, story_content: Document) -> str:
        """Save the generated Word document only."""
        output_dir = Path(self.base_path) / "data" / "outputs" / "stories"
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        try:
            filename = f"confluence_story_{timestamp}.docx"
            path = output_dir / filename
            story_content.save(path)
            print(f" Saved Word document: {path}")
            return str(path)
        except Exception as e:
            print(f"Error saving file: {e}")
            return ""
